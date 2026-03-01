"""
業務マニュアル Word ドキュメント生成スクリプト

pharma-shift シフト調整業務の全ビジネスロジックを体系化した
10章構成の業務マニュアルを python-docx で生成する。

Usage:
    pip install python-docx
    python docs/generate_manual.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "業務マニュアル_シフト調整.docx")

FONT_HEADING = "游ゴシック"
FONT_BODY = "游明朝"
COLOR_NAVY = RGBColor(0x00, 0x2B, 0x5C)
COLOR_BLUE = RGBColor(0x00, 0x52, 0x9B)
COLOR_HEADER_BG = "1F4E79"
COLOR_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ALT_ROW = "E8F0FE"


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color_hex: str):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False, color=None,
                  size=None, font_name=None):
    """Set text content and formatting for a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=COLOR_HEADER_TEXT,
                      size=Pt(10), font_name=FONT_HEADING)
        set_cell_shading(cell, COLOR_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            set_cell_text(cell, str(val), size=Pt(10), font_name=FONT_BODY)
            if row_idx % 2 == 1:
                set_cell_shading(cell, COLOR_ALT_ROW)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph("")  # spacer
    return table


def add_heading1(doc, text: str):
    """Add Chapter-level heading (Heading 1)."""
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_NAVY
        run.font.name = FONT_HEADING
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)


def add_heading2(doc, text: str):
    """Add Section-level heading (Heading 2)."""
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_BLUE
        run.font.name = FONT_HEADING
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)


def add_heading3(doc, text: str):
    """Add Subsection-level heading (Heading 3)."""
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.size = Pt(11)
        run.bold = True
        run.font.name = FONT_HEADING
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)


def add_body(doc, text: str):
    """Add body paragraph with standard formatting."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = FONT_BODY
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    return p


def add_formula(doc, text: str):
    """Add formula / code block as indented monospace paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Consolas"
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_bullet(doc, text: str, level: int = 0):
    """Add a bullet list item."""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = FONT_BODY
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    return p


def add_page_number(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)


# ---------------------------------------------------------------------------
# Chapter content functions
# ---------------------------------------------------------------------------
def chapter1(doc):
    """第1章: システム概要"""
    add_heading1(doc, "第1章　システム概要")

    add_heading2(doc, "1.1　pharma-shift とは")
    add_body(doc,
        "pharma-shift は、ツルハグループ北海道内 約62店舗の薬局を対象に、"
        "応援薬剤師（ラウンダー）の派遣・シフト調整業務を自動化する Web アプリケーションです。"
    )
    add_body(doc,
        "従来は Excel と電話で行っていたシフト調整業務を、"
        "ビジネスルールに基づく自動スコアリングと優先度管理により効率化します。"
    )

    add_heading2(doc, "1.2　4つのユーザーロール")
    add_body(doc, "本システムでは、業務上の役割に応じて以下の4つのロールを定義しています。"
                  "ユーザーは複数のロールを同時に持つことができます。")
    add_table(doc,
        ["ロール", "英語名", "主な権限"],
        [
            ["管理者", "admin", "全機能のフルアクセス。ユーザー管理、マスタ編集、全データ閲覧"],
            ["スーパーバイザー（SV）", "supervisor", "担当エリアのシフト管理、アサイン承認、HR評価"],
            ["薬局長", "store_manager", "自店舗のシフト管理、休暇申請管理、応援依頼"],
            ["ラウンダー", "rounder", "自身のシフト確認、休暇申請、アサインの承諾・辞退"],
        ],
        col_widths=[3.5, 3.0, 10.0],
    )

    add_heading2(doc, "1.3　INSERT-only ポリシー")
    add_body(doc,
        "以下のテーブルは、データの改ざんを防止するため INSERT（新規追加）のみを許可し、"
        "UPDATE（更新）・DELETE（削除）をアプリケーション層で禁止しています。"
    )
    add_bullet(doc, "audit_logs — 全システム変更の監査ログ")
    add_bullet(doc, "hr_evaluations — スタッフ評価記録")
    add_bullet(doc, "notification_logs — Zoom 通知履歴")
    add_bullet(doc, "assignment_logs — アサインメント状態遷移記録")
    add_bullet(doc, "staff_transfers — スタッフ異動記録")

    add_heading2(doc, "1.4　技術スタック")
    add_table(doc,
        ["レイヤー", "技術"],
        [
            ["バックエンド", "Django 5.1 + Django REST Framework"],
            ["認証", "django-allauth + SimpleJWT（メールベース認証）"],
            ["データベース", "PostgreSQL（開発時は SQLite）"],
            ["非同期/キュー", "Celery + Redis"],
            ["ホスティング", "Heroku"],
            ["フロントエンド", "React 18 + TypeScript + React Router"],
        ],
        col_widths=[3.5, 13.0],
    )


def chapter2(doc):
    """第2章: 店舗マスタ"""
    add_heading1(doc, "第2章　店舗マスタ")

    add_heading2(doc, "2.1　店舗の基本情報")
    add_body(doc,
        "各薬局（店舗）は以下の情報をマスタとして管理します。"
    )
    add_table(doc,
        ["フィールド", "説明", "デフォルト"],
        [
            ["name（店舗名）", "薬局の正式名称（最大100文字）", "（必須）"],
            ["area（エリア）", "地域分類（最大50文字）", "（任意）"],
            ["base_difficulty", "基本難易度（1.0〜5.0）", "3.0"],
            ["slots", "応援枠数", "1"],
            ["min_pharmacists", "1日あたり最低必要薬剤師数", "1"],
            ["monthly_working_days", "1シフト期間の営業日数（店舗独自設定）", "なし（雇用形態デフォルトを使用）"],
            ["operates_on_holidays", "祝日営業するかどうか", "False"],
            ["zoom_account", "通知先 Zoom メールアドレス", "（任意）"],
            ["is_active", "店舗の有効/無効", "True"],
        ],
        col_widths=[4.0, 8.0, 4.5],
    )

    add_heading2(doc, "2.2　初見殺しフラグ（難易度調整フラグ）")
    add_body(doc,
        "初めてその店舗に行くラウンダーにとって特に難しい要素を「初見殺しフラグ」として "
        "7つ定義しています。各フラグが ON の場合、基本難易度に加算されます。"
    )
    add_table(doc,
        ["フラグ", "日本語名", "加算値"],
        [
            ["has_controlled_medical_device", "高度管理医療機器", "+0.5"],
            ["has_toxic_substances", "毒劇物販売", "+0.5"],
            ["has_workers_comp", "労災対応", "+0.3"],
            ["has_auto_insurance", "自賠責対応", "+0.3"],
            ["has_special_public_expense", "特殊公費", "+0.4"],
            ["has_local_voucher", "地方振興券対応", "+0.2"],
            ["has_holiday_rules", "祝日出勤特殊ルール", "+0.3"],
        ],
        col_widths=[5.5, 4.5, 2.5],
    )

    add_heading2(doc, "2.3　実効難易度の算出")
    add_body(doc,
        "店舗の実効難易度（effective_difficulty）は、基本難易度に全フラグの加算値を合計し、"
        "上限5.0でキャップします。"
    )
    add_formula(doc, "effective_difficulty = min(base_difficulty + Σ(有効フラグの加算値), 5.0)")
    add_body(doc, "【計算例】")
    add_bullet(doc, "基本難易度 3.0 + 高度管理医療機器(+0.5) + 毒劇物(+0.5) = 4.0")
    add_bullet(doc, "基本難易度 4.5 + 労災(+0.3) + 自賠責(+0.3) + 特殊公費(+0.4) = min(5.5, 5.0) = 5.0（キャップ）")


def chapter3(doc):
    """第3章: スタッフ管理"""
    add_heading1(doc, "第3章　スタッフ管理")

    add_heading2(doc, "3.1　職種")
    add_table(doc,
        ["職種コード", "日本語名", "説明"],
        [
            ["pharmacist", "薬剤師", "一般薬剤師。応援・派遣の対象"],
            ["clerk", "事務員", "薬剤師人員カウントの対象外"],
            ["managing_pharmacist", "管理薬剤師", "所属店舗以外のシフトに入れない制約あり"],
        ],
        col_widths=[3.5, 3.5, 9.5],
    )

    add_heading2(doc, "3.2　雇用形態と月間所定労働日数")
    add_body(doc, "雇用形態ごとにデフォルトの月間所定労働日数が定められています。")
    add_table(doc,
        ["雇用形態", "英語名", "デフォルト日数/期間"],
        [
            ["正社員", "full_time", "22日"],
            ["パート", "part_time", "15日"],
            ["派遣", "dispatch", "20日"],
        ],
        col_widths=[3.0, 3.0, 4.0],
    )

    add_heading3(doc, "月間所定労働日数の優先順位")
    add_body(doc, "実際に適用される月間所定労働日数は、以下の優先順位で決定されます。")
    add_bullet(doc, "① スタッフ個人に設定された monthly_working_days（最優先）")
    add_bullet(doc, "② スタッフの所属店舗に設定された monthly_working_days")
    add_bullet(doc, "③ 雇用形態のデフォルト値（正社員=22、パート=15、派遣=20）")

    add_heading2(doc, "3.3　勤務ステータス（work_status）")
    add_body(doc, "スタッフの現在の勤務状態を表す4つのステータスです。")
    add_table(doc,
        ["ステータス", "日本語名", "自動アサイン対象"],
        [
            ["active", "通常勤務", "対象"],
            ["on_leave", "休職中", "対象外"],
            ["maternity", "産休・育休中", "対象外"],
            ["temporary", "臨時人員", "is_active=True の場合のみ対象"],
        ],
        col_widths=[3.0, 3.5, 5.0],
    )

    add_heading3(doc, "自動アサイン対象判定（is_auto_assignable）")
    add_body(doc, "以下の両方を満たす場合のみ、自動アサインの候補となります。")
    add_bullet(doc, "is_active = True")
    add_bullet(doc, "work_status = active")

    add_heading2(doc, "3.4　ラウンダー（応援薬剤師）")
    add_body(doc,
        "ラウンダーは Staff モデルと 1対1 で紐づく拡張モデルです。"
        "応援派遣に必要な追加属性を管理します。"
    )
    add_table(doc,
        ["フィールド", "説明", "デフォルト / 範囲"],
        [
            ["hunter_rank（HR値）", "ハンターランク。スキルレベルを示す数値", "20.0（下限0.0）"],
            ["can_work_alone", "一人薬剤師（ソロ勤務）が可能か", "False"],
            ["max_prescriptions", "1日の処理可能処方箋枚数", "30（範囲: 1〜50）"],
            ["has_car", "自家用車の有無", "False"],
            ["can_long_distance", "遠距離（他エリア）への派遣可否", "False"],
            ["managing_pharmacist_years", "管理薬剤師経験年数", "0"],
        ],
        col_widths=[4.5, 6.5, 5.5],
    )

    add_heading3(doc, "初期HR値の算出")
    add_body(doc, "ラウンダー登録時、管理薬剤師経験年数から初期HR値を算出します。")
    add_formula(doc, "initial_hr = min(managing_pharmacist_years × 5, 30)")
    add_bullet(doc, "例: 3年 → min(15, 30) = 15 HR")
    add_bullet(doc, "例: 7年 → min(35, 30) = 30 HR（キャップ）")

    add_heading2(doc, "3.5　店舗経験（RounderStoreExperience）")
    add_body(doc,
        "ラウンダーが過去に訪問した店舗の経験を記録します。"
        "自動スコアリングで経験者ボーナスの判定に使用されます。"
    )
    add_bullet(doc, "first_visit_date — 初回訪問日")
    add_bullet(doc, "last_visit_date — 最終訪問日")
    add_bullet(doc, "visit_count — 累計訪問回数")

    add_heading2(doc, "3.6　ラウンダー不稼働期間（RounderUnavailability）")
    add_body(doc,
        "特定のシフト期間中、ラウンダーを応援アサインの対象外とする設定です。"
        "長期研修や個人的な事情で一定期間派遣不可の場合に使用します。"
        "ラウンダーとシフト期間の組み合わせでユニーク制約があります。"
    )

    add_heading2(doc, "3.7　スタッフ異動（StaffTransfer）")
    add_body(doc,
        "スタッフの店舗間異動を記録します。INSERT-only（変更・削除不可）の監査テーブルです。"
    )
    add_bullet(doc, "from_store — 異動元店舗")
    add_bullet(doc, "to_store — 異動先店舗")
    add_bullet(doc, "reason — 異動理由")
    add_bullet(doc, "transferred_by — 異動を実行したユーザー")


def chapter4(doc):
    """第4章: シフトサイクルとルール"""
    add_heading1(doc, "第4章　シフトサイクルとルール")

    add_heading2(doc, "4.1　シフト期間")
    add_body(doc,
        "本システムのシフトは月の16日から翌月15日までを1サイクルとします。"
    )
    add_table(doc,
        ["項目", "内容"],
        [
            ["開始日", "毎月16日"],
            ["終了日", "翌月15日"],
            ["期間", "約30日間"],
            ["休暇申請締切", "シフト開始日の15日前"],
            ["確定フラグ", "is_finalized = True でスケジュール確定（ロック）"],
        ],
        col_widths=[4.0, 12.5],
    )
    add_body(doc, "【例】2026年4月シフト期間: 4月16日〜5月15日、休暇申請締切: 4月1日")

    add_heading2(doc, "4.2　シフト種別")
    add_table(doc,
        ["種別", "日本語名", "説明"],
        [
            ["full", "全日", "終日勤務"],
            ["morning", "午前", "午前のみ勤務"],
            ["afternoon", "午後", "午後のみ勤務"],
        ],
        col_widths=[3.0, 3.0, 10.5],
    )

    add_heading2(doc, "4.3　休暇種別")
    add_table(doc,
        ["種別", "日本語名", "説明"],
        [
            ["paid", "有給", "有給休暇"],
            ["holiday", "公休", "所定休日"],
            ["sick", "病欠", "傷病による休暇"],
            ["other", "その他", "その他の休暇"],
        ],
        col_widths=[3.0, 3.0, 10.5],
    )

    add_heading2(doc, "4.4　ダブルブッキング防止ルール")
    add_body(doc,
        "同一スタッフ・同一日のシフト重複を以下のルールで防止します。"
    )
    add_table(doc,
        ["ルール", "既存シフト", "新規シフト", "結果"],
        [
            ["ルール1", "全日（full）", "任意のシフト", "エラー（重複不可）"],
            ["ルール2", "任意のシフト", "全日（full）", "エラー（重複不可）"],
            ["ルール3", "午前（morning）", "午後（afternoon）", "許可（半日ずつ可）"],
            ["ルール4", "午前（morning）", "午前（morning）", "エラー（同じ時間帯の重複不可）"],
        ],
        col_widths=[2.5, 3.5, 3.5, 7.0],
    )

    add_heading2(doc, "4.5　管理薬剤師制約")
    add_body(doc,
        "管理薬剤師（managing_pharmacist）は、法令上の責任から "
        "所属店舗以外のシフトに入ることができません。"
    )
    add_body(doc,
        "もし所属店舗以外の店舗にシフトを登録しようとした場合、"
        "「管理薬剤師（氏名）は所属店舗以外のシフトに入れません」というエラーが表示されます。"
    )

    add_heading2(doc, "4.6　最低人員チェック")
    add_body(doc,
        "店舗には最低必要薬剤師数（min_pharmacists）が設定されています。"
        "薬剤師が休暇を取るなどして店舗を離れる場合、残る薬剤師数が最低人員を下回らないか検証します。"
    )
    add_bullet(doc, "チェック対象: 所属店舗が設定されている薬剤師（事務員は対象外）")
    add_bullet(doc, "判定方法: シフト未登録のスタッフは所属店舗にいると見なす")
    add_bullet(doc, "カウント条件: is_active=True かつ work_status=active のスタッフのみ")

    add_heading2(doc, "4.7　月間上限チェック")
    add_body(doc,
        "スタッフの月間所定労働日数を超えるシフト登録を防止します。"
    )
    add_bullet(doc, "休暇シフトは労働日数にカウントしない")
    add_bullet(doc, "午前＋午後を同一日に登録した場合は1日としてカウント")
    add_bullet(doc, "上限値 = スタッフの effective_monthly_working_days（3.2節の優先順位に従う）")


def chapter5(doc):
    """第5章: 応援枠（クエスト）システム"""
    add_heading1(doc, "第5章　応援枠（クエスト）システム")

    add_heading2(doc, "5.1　応援枠とは")
    add_body(doc,
        "応援枠（Support Slot）は、「この店舗にこの日、応援薬剤師が必要」という1件の派遣要求です。"
        "ゲームのクエストボードのように、優先度と難易度が設定され、"
        "適切なラウンダーが自動的に候補として提案されます。"
    )

    add_heading2(doc, "5.2　優先度（P1〜P5）")
    add_body(doc, "応援枠は緊急度に応じて5段階の優先度が設定されます。")
    add_table(doc,
        ["優先度", "トリガー", "対応レベル"],
        [
            ["P1", "緊急欠員・法的コンプライアンスリスク", "最優先。即日対応必須"],
            ["P2", "義務有給5日消化の期限迫る", "高優先。期限内に必ず対応"],
            ["P3", "薬局長・管理薬剤師の公休確保", "中優先。計画的に対応"],
            ["P4", "希望休・健康診断", "通常。スケジュール調整で対応"],
            ["P5", "その他有給・任意休暇", "低優先。空きがあれば対応"],
        ],
        col_widths=[2.0, 6.5, 8.0],
    )

    add_heading2(doc, "5.3　処方予測グレード")
    add_body(doc,
        "応援日の処方箋枚数予測をA〜Eの5段階で評価し、"
        "難易度計算に反映します。"
    )
    add_table(doc,
        ["グレード", "意味", "難易度ペナルティ（HR単位）"],
        [
            ["A", "多い", "+10"],
            ["B", "やや多い", "+5"],
            ["C", "普通", "±0"],
            ["D", "やや少ない", "-5"],
            ["E", "少ない", "-10"],
        ],
        col_widths=[2.5, 4.0, 6.0],
    )

    add_heading2(doc, "5.4　実効難易度（HR単位）の算出")
    add_body(doc,
        "応援枠の実効難易度は、店舗の基本難易度をベースに、"
        "当日の状況を加味してHR単位で算出されます。"
    )
    add_formula(doc, "effective_difficulty_hr = max(")
    add_formula(doc, "    base_difficulty × 10")
    add_formula(doc, "    - (5  ※薬局長在席の場合)")
    add_formula(doc, "    - (出勤薬剤師数 × 3)")
    add_formula(doc, "    + (ソロ勤務時間 × 2)")
    add_formula(doc, "    + 処方予測ペナルティ,")
    add_formula(doc, "    0  ※下限0")
    add_formula(doc, ")")
    add_body(doc, "")
    add_body(doc, "【計算例1】基本難易度 3.0、薬局長在席、薬剤師1名出勤、ソロなし、予測C:")
    add_formula(doc, "30 - 5 - 3 + 0 + 0 = 22 HR")
    add_body(doc, "【計算例2】基本難易度 2.0、薬局長不在、薬剤師0名、ソロ2時間、予測B:")
    add_formula(doc, "20 - 0 - 0 + 4 + 5 = 29 HR")

    add_heading2(doc, "5.5　応援枠の主要フィールド")
    add_table(doc,
        ["フィールド", "説明"],
        [
            ["store", "応援が必要な店舗"],
            ["date", "応援日"],
            ["shift_period", "所属するシフト期間"],
            ["priority", "優先度（P1〜P5）"],
            ["base_difficulty", "基本難易度（店舗設定を上書き可能）"],
            ["attending_pharmacists", "当日出勤する薬剤師数"],
            ["attending_clerks", "当日出勤する事務員数"],
            ["has_chief_present", "薬局長が在席するか"],
            ["solo_hours", "ソロ勤務が必要な時間数"],
            ["prescription_forecast", "処方予測グレード（A〜E）"],
            ["effective_difficulty_hr", "算出された実効難易度（HR単位）"],
            ["required_hr", "必要最低HR値（デフォルト=実効難易度）"],
        ],
        col_widths=[5.0, 11.5],
    )


def chapter6(doc):
    """第6章: 自動アサインとスコアリング"""
    add_heading1(doc, "第6章　自動アサインとスコアリング")

    add_heading2(doc, "6.1　候補者の自動生成")
    add_body(doc,
        "応援枠に対して、システムが自動的に最適な候補者を提案します。"
        "デフォルトで上位5名の候補者リストを生成します。"
    )

    add_heading3(doc, "候補者フィルタ（3条件）")
    add_body(doc, "以下の3条件を全て満たすスタッフのみが候補者の対象となります。")
    add_bullet(doc, "① is_active = True（有効なスタッフ）")
    add_bullet(doc, "② is_rounder = True（ラウンダーとして登録済み）")
    add_bullet(doc, "③ work_status = active（通常勤務中）")

    add_heading2(doc, "6.2　前提条件チェック（6項目）")
    add_body(doc,
        "フィルタを通過した候補者に対し、さらに6つの前提条件をチェックします。"
        "1つでも不適合があればその候補者はスキップされます。"
    )
    add_table(doc,
        ["#", "チェック項目", "エラー条件"],
        [
            ["1", "不稼働期間チェック", "該当シフト期間に不稼働登録がある場合"],
            ["2", "HR値チェック", "rounder.hunter_rank < slot.required_hr"],
            ["3", "シフト競合チェック", "応援日に既存シフト（勤務・休暇）が登録済み"],
            ["4", "アサイン競合チェック", "同日に他の応援枠の候補 or 確定アサインが存在"],
            ["5", "ソロ勤務能力", "slot.solo_hours > 0 だが can_work_alone = False"],
            ["6", "遠距離派遣能力", "異なるエリアだが can_long_distance = False"],
        ],
        col_widths=[1.0, 4.0, 11.5],
    )

    add_heading2(doc, "6.3　スコアリング（4要素）")
    add_body(doc,
        "前提条件を全てクリアした候補者に対し、以下の4要素でスコアを算出します。"
        "スコアが高いほど優先的に推薦されます。"
    )
    add_table(doc,
        ["要素", "スコア", "条件"],
        [
            ["店舗経験ボーナス", "+100", "過去にその店舗を訪問した経験がある"],
            ["直近訪問ボーナス", "+50", "最終訪問日が90日以内"],
            ["HR余裕度ボーナス", "最大+20", "min((HR値 - 必要HR) × 2, 20)"],
            ["同エリアボーナス", "+15", "ラウンダーの所属店舗と応援先が同じエリア"],
        ],
        col_widths=[4.0, 2.5, 10.0],
    )
    add_body(doc, "")
    add_body(doc, "【スコア計算例】")
    add_bullet(doc, "店舗経験あり、90日以内訪問、HR50で必要HR35、同エリア:")
    add_formula(doc, "100 + 50 + min((50-35)×2, 20) + 15 = 185点")
    add_bullet(doc, "未経験、HR40で必要HR40、異エリア:")
    add_formula(doc, "0 + 0 + min(0×2, 20) + 0 = 0点")

    add_heading2(doc, "6.4　エリア判定")
    add_body(doc,
        "ラウンダーの所属店舗のエリアと応援先店舗のエリアを比較します。"
        "どちらかのエリアが未設定（空白）の場合は同一エリアとして扱います。"
    )

    add_heading2(doc, "6.5　アサインのステータス遷移")
    add_body(doc, "各アサインは以下の5つのステータスで管理されます。")
    add_table(doc,
        ["ステータス", "日本語", "説明"],
        [
            ["candidate", "候補", "自動生成された初期状態"],
            ["confirmed", "確定", "SV・薬局長により承認された状態"],
            ["rejected", "辞退", "ラウンダーまたは管理者が明示的に辞退"],
            ["cancelled", "キャンセル", "確定後にキャンセルされた状態"],
            ["handed_over", "引継", "他のラウンダーに引き継がれた状態"],
        ],
        col_widths=[3.0, 2.5, 11.0],
    )
    add_body(doc,
        "全てのステータス遷移は AssignmentLog に INSERT-only で記録され、"
        "誰がいつ変更したかの監査証跡が残ります。"
    )


def chapter7(doc):
    """第7章: HR（ハンターランク）システム"""
    add_heading1(doc, "第7章　HR（ハンターランク）システム")

    add_heading2(doc, "7.1　HRシステムの目的")
    add_body(doc,
        "HR（ハンターランク）は、ラウンダーのスキルレベルを数値化する独自の評価指標です。"
        "応援枠の難易度要件と照合することで、"
        "適切なスキルレベルのラウンダーを自動的にマッチングします。"
    )

    add_heading2(doc, "7.2　成長曲線（3段階）")
    add_body(doc,
        "累積評価ポイントからHR値への変換は、3段階の非線形成長曲線に従います。"
        "経験が浅いうちは急速に成長し、ベテランになるほど成長が緩やかになる設計です。"
    )
    add_table(doc,
        ["ポイント範囲", "換算レート", "HR値の範囲"],
        [
            ["0〜30ポイント", "1ポイント = 2 HR", "0〜60 HR"],
            ["31〜60ポイント", "1ポイント = 1 HR", "61〜90 HR"],
            ["61ポイント以上", "1ポイント = 0.5 HR", "91 HR〜"],
        ],
        col_widths=[4.0, 4.5, 4.0],
    )

    add_body(doc, "")
    add_body(doc, "【換算式】")
    add_formula(doc, "points ≤ 30  → HR = points × 2")
    add_formula(doc, "points ≤ 60  → HR = 60 + (points - 30) × 1")
    add_formula(doc, "points > 60  → HR = 90 + (points - 60) × 0.5")

    add_body(doc, "")
    add_body(doc, "【計算例】")
    add_bullet(doc, "10ポイント → 20 HR")
    add_bullet(doc, "30ポイント → 60 HR")
    add_bullet(doc, "45ポイント → 60 + 15 = 75 HR")
    add_bullet(doc, "60ポイント → 90 HR")
    add_bullet(doc, "100ポイント → 90 + 20 = 110 HR")

    add_heading2(doc, "7.3　評価の種類")
    add_table(doc,
        ["評価種別", "評価者", "スコア範囲", "説明"],
        [
            ["supervisor（SV評価）", "スーパーバイザー", "-1.0 〜 +1.0", "応援業務の品質を評価"],
            ["self（自己評価）", "ラウンダー本人", "-0.5 〜 +0.5", "自己振り返り評価（レンジ制限あり）"],
        ],
        col_widths=[4.0, 3.5, 3.0, 6.0],
    )

    add_heading2(doc, "7.4　期間サマリーの算出")
    add_body(doc, "シフト期間ごとにHR値を再計算します。")
    add_bullet(doc, "① SV評価の合計を算出 → supervisor_total")
    add_bullet(doc, "② 自己評価の合計を算出 → self_total")
    add_bullet(doc, "③ 繰越ポイントを決定:")
    add_bullet(doc, "初回期間: min(管理薬剤師経験年数 × 5, 30)", level=1)
    add_bullet(doc, "2回目以降: 前期のtotal_points × 0.7（繰越係数）", level=1)
    add_bullet(doc, "④ 合計: total_points = 繰越 + supervisor_total + self_total")
    add_bullet(doc, "⑤ 成長曲線を適用: computed_hr = points_to_hr(max(total_points, 0))")
    add_bullet(doc, "⑥ ラウンダーの hunter_rank を更新")

    add_heading2(doc, "7.5　公平性チェック")

    add_heading3(doc, "連続ネガティブ検知")
    add_body(doc,
        "同一評価者が同一ラウンダーに対して直近2期連続で -1.0 以下のスコアを付けた場合、"
        "新しい評価に requires_approval = True フラグが設定され、管理者の承認が必要になります。"
    )

    add_heading3(doc, "バイアス検知")
    add_body(doc, "評価者の偏りを統計的に検出します。")
    add_bullet(doc, "全評価者の -1.0 評価比率（グローバル平均）を算出")
    add_bullet(doc, "該当評価者の -1.0 評価比率を算出")
    add_bullet(doc, "評価件数が3件以上 かつ 個人比率がグローバル平均の2倍を超える場合 → アラート発報")


def chapter8(doc):
    """第8章: 休暇管理"""
    add_heading1(doc, "第8章　休暇管理")

    add_heading2(doc, "8.1　義務有給5日消化")
    add_body(doc,
        "労働基準法に基づき、年5日の有給休暇取得が義務付けられています。"
        "本システムでは、各スタッフの有給消化期限を管理し、自動でアラートを発報します。"
    )

    add_heading3(doc, "有給消化期限")
    add_body(doc, "スタッフには以下の2つの期限カテゴリがあります。")
    add_table(doc,
        ["カテゴリ", "期限日", "対象"],
        [
            ["デフォルト", "9月15日", "大多数のスタッフ"],
            ["代替", "2月15日", "入社時期等による特例"],
        ],
        col_widths=[3.0, 3.0, 10.5],
    )

    add_heading2(doc, "8.2　アラート3段階")
    add_body(doc, "有給消化状況に応じて、3段階のアラートが発報されます。")
    add_table(doc,
        ["レベル", "条件", "メッセージ", "対応"],
        [
            ["overdue（期限超過）", "期限日を過ぎても未消化", "義務有給期限超過（残X日未消化）", "至急対応"],
            ["urgent（緊急）", "期限まで7日以内", "義務有給期限7日以内（残X日未消化）", "優先対応"],
            ["warning（警告）", "期限まで30日以内", "義務有給期限30日以内（残X日未消化）", "計画的に対応"],
        ],
        col_widths=[3.5, 3.5, 5.5, 4.0],
    )

    add_heading2(doc, "8.3　休暇申請ワークフロー")
    add_body(doc, "休暇申請は以下のフローで処理されます。")
    add_bullet(doc, "① ラウンダー / 薬局長がシステム上で休暇申請を作成")
    add_bullet(doc, "② シフト期間の休暇申請締切（開始15日前）までに申請")
    add_bullet(doc, "③ 締切後の申請は is_late_request = True フラグが付与される")
    add_bullet(doc, "④ SV / 管理者が申請を review（承認 or 却下）")
    add_bullet(doc, "⑤ 承認された休暇はシフトに自動反映")


def chapter9(doc):
    """第9章: スタッフィングチェック"""
    add_heading1(doc, "第9章　スタッフィングチェック")

    add_heading2(doc, "9.1　概要")
    add_body(doc,
        "スタッフィングチェックは、各店舗が十分な人員を確保できているかを "
        "日次で検証する仕組みです。週間スケジュールテンプレートと "
        "日別オーバーライドの2層構造で運用します。"
    )

    add_heading2(doc, "9.2　週間スケジュール（StoreWeeklySchedule）")
    add_body(doc,
        "各店舗の曜日ごとの基本運営設定を定義するテンプレートです。"
        "店舗×曜日（0=月曜〜6=日曜）の組み合わせでユニーク制約があります。"
    )
    add_table(doc,
        ["フィールド", "説明", "範囲"],
        [
            ["is_open", "その曜日に営業するか", "True / False"],
            ["open_time / close_time", "営業時間", "時刻"],
            ["staffing_delta", "人員調整係数", "-5.0 〜 +5.0"],
            ["note", "備考", "自由記述"],
        ],
        col_widths=[4.5, 6.5, 5.5],
    )
    add_body(doc,
        "一括更新エンドポイント（/staffing/weekly_schedules/bulk_upsert/）により、"
        "7曜日分をまとめてアトミックに登録・更新できます。"
    )

    add_heading2(doc, "9.3　日別オーバーライド（DailyScheduleOverride）")
    add_body(doc,
        "週間テンプレートの例外を日単位で設定します。"
        "臨時休業、当番薬局、特別営業日などに使用します。"
    )
    add_bullet(doc, "is_open = False → 臨時休業")
    add_bullet(doc, "is_open = True → 通常休業日だが特別に営業（当番薬局 等）")
    add_bullet(doc, "note に理由を記載（例: 「当番薬局」「臨時休業」）")

    add_heading2(doc, "9.4　スタッフィング調整（StaffingAdjustment）")
    add_body(doc,
        "人員の過不足を予測・調整するための記録です。"
        "手動入力と機械学習モデルによる予測の2つのソースで管理します。"
    )
    add_table(doc,
        ["ソース", "説明", "入力者"],
        [
            ["manual", "SVが手動で設定する調整値", "スーパーバイザー"],
            ["model", "MLモデルが日次で予測する調整値", "自動（Celeryタスク）"],
        ],
        col_widths=[3.0, 8.5, 5.0],
    )
    add_body(doc,
        "店舗 × シフト期間 × 日付 × ソース の組み合わせでユニーク制約があり、"
        "手動とモデルの予測値を独立して追跡できます。"
    )


def chapter10(doc):
    """第10章: 通知と自動化"""
    add_heading1(doc, "第10章　通知と自動化")

    add_heading2(doc, "10.1　Zoom Chat 連携")
    add_body(doc,
        "本システムは Zoom の Server-to-Server OAuth を利用して、"
        "シフト関連の通知を Zoom Chat 経由で自動送信します。"
    )

    add_heading3(doc, "認証設定")
    add_body(doc, "以下の3つの環境変数が必要です。")
    add_bullet(doc, "ZOOM_ACCOUNT_ID — Zoom アカウント ID")
    add_bullet(doc, "ZOOM_CLIENT_ID — OAuth クライアント ID")
    add_bullet(doc, "ZOOM_CLIENT_SECRET — OAuth クライアントシークレット")

    add_heading3(doc, "通知の仕組み")
    add_bullet(doc, "店舗の zoom_account（メールアドレス）宛に Zoom Chat メッセージを送信")
    add_bullet(doc, "送信結果は NotificationLog に INSERT-only で記録")
    add_bullet(doc, "送信失敗時はエラーメッセージが error_message フィールドに記録される")
    add_bullet(doc, "zoom_account が未設定の店舗には通知されない")

    add_heading2(doc, "10.2　Celery 定期タスク")
    add_body(doc,
        "以下の定期タスクが Celery Beat により自動実行されます。"
        "全ての時刻は日本標準時（JST）です。"
    )
    add_table(doc,
        ["タスク名", "実行時刻", "内容"],
        [
            ["fetch_weather_daily", "毎日 05:30", "天候予報データの取得"],
            ["fetch_idwr_weekly", "毎週火曜 05:00", "感染症発生動向データ（IDWR）の取得"],
            ["scrape_musubi_prescriptions", "毎日 06:00", "POSから処方箋実績データを取得"],
            ["generate_prescription_forecasts", "毎週月曜 07:00", "MLモデルによる処方箋枚数予測の生成"],
            ["check_paid_leave_alerts", "毎日 08:00", "義務有給消化アラートの発報"],
            ["check_leave_request_deadline", "毎日 09:00", "休暇申請締切のリマインド（3日前通知）"],
            ["alert_unfilled_slots", "毎日 10:00", "P1/P2の未充足応援枠をSVにアラート"],
            ["check_all_evaluator_bias", "毎日 11:00", "HR評価のバイアス検知"],
        ],
        col_widths=[5.5, 3.5, 7.5],
    )

    add_heading2(doc, "10.3　Celery の起動方法")
    add_body(doc, "ワーカーとビートスケジューラーを同時に起動します。")
    add_formula(doc, "celery -A config worker --beat --loglevel=info")
    add_body(doc, "※ backend/ ディレクトリから実行してください。")


# ---------------------------------------------------------------------------
# Appendix: Key constants
# ---------------------------------------------------------------------------
def appendix(doc):
    """付録: 主要定数・閾値一覧"""
    add_heading1(doc, "付録　主要定数・閾値一覧")

    add_body(doc, "システム全体で使用される主要な定数と閾値の一覧です。")

    add_heading2(doc, "A.1　店舗関連")
    add_table(doc,
        ["パラメータ", "値", "説明"],
        [
            ["難易度上限", "5.0", "effective_difficulty の最大値"],
            ["基本難易度範囲", "1.0〜5.0", "店舗ごとの基本難易度"],
            ["最低薬剤師数デフォルト", "1", "min_pharmacists のデフォルト"],
        ],
        col_widths=[4.5, 3.0, 9.0],
    )

    add_heading2(doc, "A.2　ラウンダー関連")
    add_table(doc,
        ["パラメータ", "値", "説明"],
        [
            ["HR初期値", "20.0", "hunter_rank のデフォルト"],
            ["HR下限", "0.0", "hunter_rank の最小値"],
            ["初期HRキャップ", "30", "管理薬剤師経験からの初期HR上限"],
            ["年数→HR係数", "×5", "管理薬剤師経験年数の変換係数"],
            ["処方箋上限/日", "50", "max_prescriptions の最大値"],
            ["処方箋デフォルト/日", "30", "max_prescriptions のデフォルト"],
        ],
        col_widths=[4.5, 3.0, 9.0],
    )

    add_heading2(doc, "A.3　HR評価関連")
    add_table(doc,
        ["パラメータ", "値", "説明"],
        [
            ["繰越係数", "0.7（70%）", "前期ポイントの繰越率"],
            ["評価スコア範囲", "-1.0〜+1.0", "SV評価のスコア範囲"],
            ["自己評価スコア範囲", "-0.5〜+0.5", "自己評価のスコア範囲"],
            ["バイアス検知サンプル数", "3件以上", "検知の最低評価件数"],
            ["バイアス検知閾値", "2×グローバル平均", "アラート発報の比率閾値"],
            ["連続ネガティブ検知", "2期連続", "承認必須フラグの条件"],
        ],
        col_widths=[4.5, 3.5, 8.5],
    )

    add_heading2(doc, "A.4　スコアリング関連")
    add_table(doc,
        ["パラメータ", "値", "説明"],
        [
            ["店舗経験ボーナス", "+100", "過去訪問経験がある場合"],
            ["直近訪問ボーナス", "+50", "90日以内の訪問がある場合"],
            ["HR余裕度ボーナス上限", "+20", "min(margin×2, 20)"],
            ["同エリアボーナス", "+15", "所属と応援先が同エリア"],
            ["候補者リスト上限", "5名", "デフォルトの候補者数"],
            ["直近訪問ウィンドウ", "90日", "直近訪問ボーナスの判定期間"],
        ],
        col_widths=[4.5, 3.0, 9.0],
    )

    add_heading2(doc, "A.5　実効難易度関連")
    add_table(doc,
        ["パラメータ", "値", "説明"],
        [
            ["薬局長在席ディスカウント", "-5 HR", "has_chief_present = True 時"],
            ["薬剤師ディスカウント", "-3 HR/人", "出勤薬剤師1人あたり"],
            ["ソロ勤務ペナルティ", "+2 HR/時間", "ソロ勤務1時間あたり"],
            ["処方予測A", "+10 HR", "処方箋多い"],
            ["処方予測B", "+5 HR", "処方箋やや多い"],
            ["処方予測C", "±0 HR", "普通"],
            ["処方予測D", "-5 HR", "処方箋やや少ない"],
            ["処方予測E", "-10 HR", "処方箋少ない"],
        ],
        col_widths=[4.5, 3.0, 9.0],
    )

    add_heading2(doc, "A.6　労働日数関連")
    add_table(doc,
        ["パラメータ", "値", "説明"],
        [
            ["正社員デフォルト", "22日/期間", "full_time の月間所定労働日数"],
            ["パートデフォルト", "15日/期間", "part_time の月間所定労働日数"],
            ["派遣デフォルト", "20日/期間", "dispatch の月間所定労働日数"],
        ],
        col_widths=[4.5, 3.0, 9.0],
    )

    add_heading2(doc, "A.7　有給管理関連")
    add_table(doc,
        ["パラメータ", "値", "説明"],
        [
            ["義務有給日数", "5日/期間", "年間最低取得日数"],
            ["警告閾値", "30日前", "warning アラート発報"],
            ["緊急閾値", "7日前", "urgent アラート発報"],
            ["期限超過閾値", "0日（期限経過）", "overdue アラート発報"],
        ],
        col_widths=[4.5, 3.0, 9.0],
    )


# ---------------------------------------------------------------------------
# Main: Build document
# ---------------------------------------------------------------------------
def build_document():
    """Build the complete manual document."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # ---- Title page ----
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("業務マニュアル")
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_NAVY
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("シフト調整業務")
    run.font.size = Pt(20)
    run.font.color.rgb = COLOR_BLUE
    run.font.name = FONT_HEADING
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)

    doc.add_paragraph("")

    system_name = doc.add_paragraph()
    system_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = system_name.add_run("pharma-shift ラウンダー派遣管理システム")
    run.font.size = Pt(14)
    run.font.name = FONT_HEADING
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)

    for _ in range(8):
        doc.add_paragraph("")

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run("ツルハグループ 薬局事業部")
    run.font.size = Pt(12)
    run.font.name = FONT_HEADING
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)

    # ---- Page break before TOC ----
    doc.add_page_break()

    # ---- Table of Contents header ----
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目 次")
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_NAVY
    run.bold = True
    run.font.name = FONT_HEADING
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)

    doc.add_paragraph("")

    toc_entries = [
        ("第1章", "システム概要"),
        ("第2章", "店舗マスタ"),
        ("第3章", "スタッフ管理"),
        ("第4章", "シフトサイクルとルール"),
        ("第5章", "応援枠（クエスト）システム"),
        ("第6章", "自動アサインとスコアリング"),
        ("第7章", "HR（ハンターランク）システム"),
        ("第8章", "休暇管理"),
        ("第9章", "スタッフィングチェック"),
        ("第10章", "通知と自動化"),
        ("付録", "主要定数・閾値一覧"),
    ]

    for chapter_num, chapter_title in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(f"{chapter_num}　{chapter_title}")
        run.font.size = Pt(12)
        run.font.name = FONT_HEADING
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        p.paragraph_format.space_after = Pt(6)

    # ---- Chapters ----
    doc.add_page_break()
    chapter1(doc)

    doc.add_page_break()
    chapter2(doc)

    doc.add_page_break()
    chapter3(doc)

    doc.add_page_break()
    chapter4(doc)

    doc.add_page_break()
    chapter5(doc)

    doc.add_page_break()
    chapter6(doc)

    doc.add_page_break()
    chapter7(doc)

    doc.add_page_break()
    chapter8(doc)

    doc.add_page_break()
    chapter9(doc)

    doc.add_page_break()
    chapter10(doc)

    doc.add_page_break()
    appendix(doc)

    # ---- Page numbers ----
    add_page_number(doc)

    # ---- Save ----
    doc.save(OUTPUT_FILE)
    print(f"Generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
